"""
Report generator module for creating various output formats.
Generates JSON, DOCX, and PDF reports from analysis results.
"""
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.config import settings
from app.schemas import AnalysisReport

logging.basicConfig(level=settings.log_level)
logger = logging.getLogger(__name__)


class ReportGenerator:
    """Generates reports in various formats from analysis results."""
    
    def __init__(self, report: AnalysisReport):
        """
        Initialize report generator.
        
        Args:
            report: AnalysisReport object to generate reports from
        """
        self.report = report
        self.report_dir = Path(settings.report_dir)
        self.report_dir.mkdir(exist_ok=True)
        
    def generate_json(self, filename: Optional[str] = None) -> str:
        """
        Generate JSON report.
        
        Args:
            filename: Optional custom filename
            
        Returns:
            Path to generated JSON file
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"report_{timestamp}.json"
        
        filepath = self.report_dir / filename
        
        # Convert report to dict
        report_dict = self.report.model_dump()
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(report_dict, f, ensure_ascii=False, indent=2)
        
        logger.info(f"JSON report generated: {filepath}")
        return str(filepath)
    
    def generate_docx(self, filename: Optional[str] = None) -> str:
        """
        Generate DOCX report.
        
        Args:
            filename: Optional custom filename
            
        Returns:
            Path to generated DOCX file
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"report_{timestamp}.docx"
        
        filepath = self.report_dir / filename
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Отчёт анализа документа', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Document info
        doc.add_heading('Информация о документе', 1)
        doc.add_paragraph(f'Документ: {self.report.document_name}')
        doc.add_paragraph(f'Всего частей: {self.report.total_chunks}')
        doc.add_paragraph(f'Найдено проблем: {self.report.total_findings}')
        doc.add_paragraph(f'Дата анализа: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Global analysis
        if self.report.global_analysis:
            doc.add_heading('Глобальный анализ', 1)
            
            doc.add_heading('Общее резюме', 2)
            doc.add_paragraph(self.report.global_analysis.overall_summary)
            
            if self.report.global_analysis.document_structure_issues:
                doc.add_heading('Проблемы структуры', 2)
                for issue in self.report.global_analysis.document_structure_issues:
                    doc.add_paragraph(issue, style='List Bullet')
            
            if self.report.global_analysis.consistency_issues:
                doc.add_heading('Проблемы согласованности', 2)
                for issue in self.report.global_analysis.consistency_issues:
                    doc.add_paragraph(issue, style='List Bullet')
            
            if self.report.global_analysis.recommendations:
                doc.add_heading('Рекомендации', 2)
                for rec in self.report.global_analysis.recommendations:
                    doc.add_paragraph(rec, style='List Bullet')
        
        # Summary by severity
        doc.add_heading('Статистика по серьёзности', 1)
        high = len(self.report.get_findings_by_severity('high'))
        medium = len(self.report.get_findings_by_severity('medium'))
        low = len(self.report.get_findings_by_severity('low'))
        
        doc.add_paragraph(f'Высокая: {high}')
        doc.add_paragraph(f'Средняя: {medium}')
        doc.add_paragraph(f'Низкая: {low}')
        
        # Detailed findings
        doc.add_heading('Детальные результаты по частям', 1)
        
        for chunk_analysis in self.report.chunk_analyses:
            doc.add_heading(f'Часть {chunk_analysis.chunk_id}', 2)
            
            doc.add_heading('Краткое содержание:', 3)
            doc.add_paragraph(chunk_analysis.summary)
            
            if chunk_analysis.findings:
                doc.add_heading(f'Найдено проблем: {len(chunk_analysis.findings)}', 3)
                
                for i, finding in enumerate(chunk_analysis.findings, 1):
                    p = doc.add_paragraph()
                    p.add_run(f'Проблема {i}: ').bold = True
                    
                    # Severity color
                    severity_run = p.add_run(f'[{finding.severity.upper()}] ')
                    if finding.severity == 'high':
                        severity_run.font.color.rgb = RGBColor(255, 0, 0)
                    elif finding.severity == 'medium':
                        severity_run.font.color.rgb = RGBColor(255, 165, 0)
                    else:
                        severity_run.font.color.rgb = RGBColor(128, 128, 128)
                    
                    p.add_run(f'{finding.issue_type}')
                    
                    doc.add_paragraph(f'Цитата: "{finding.quote}"', style='List Bullet 2')
                    doc.add_paragraph(f'Проблема: {finding.problem}', style='List Bullet 2')
                    doc.add_paragraph(f'Рекомендация: {finding.recommendation}', style='List Bullet 2')
            else:
                doc.add_paragraph('Проблем не найдено.')
        
        doc.save(filepath)
        logger.info(f"DOCX report generated: {filepath}")
        return str(filepath)
    
    def generate_html(self, filename: Optional[str] = None) -> str:
        """
        Generate HTML report.
        
        Args:
            filename: Optional custom filename
            
        Returns:
            Path to generated HTML file
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"report_{timestamp}.html"
        
        filepath = self.report_dir / filename
        
        html_content = self._build_html()
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        logger.info(f"HTML report generated: {filepath}")
        return str(filepath)
    
    def _build_html(self) -> str:
        """Build HTML content for report."""
        high = len(self.report.get_findings_by_severity('high'))
        medium = len(self.report.get_findings_by_severity('medium'))
        low = len(self.report.get_findings_by_severity('low'))
        
        html = f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Отчёт анализа документа</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        .header {{
            background-color: #2c3e50;
            color: white;
            padding: 20px;
            border-radius: 5px;
            margin-bottom: 20px;
        }}
        .section {{
            background-color: white;
            padding: 20px;
            margin-bottom: 20px;
            border-radius: 5px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        .severity-high {{ color: #e74c3c; font-weight: bold; }}
        .severity-medium {{ color: #f39c12; font-weight: bold; }}
        .severity-low {{ color: #95a5a6; font-weight: bold; }}
        .finding {{
            border-left: 4px solid #3498db;
            padding-left: 15px;
            margin: 15px 0;
        }}
        .quote {{
            background-color: #ecf0f1;
            padding: 10px;
            border-radius: 3px;
            font-style: italic;
            margin: 10px 0;
        }}
        h1, h2, h3 {{ color: #2c3e50; }}
        .stats {{
            display: flex;
            justify-content: space-around;
            margin: 20px 0;
        }}
        .stat-box {{
            text-align: center;
            padding: 15px;
            background-color: #ecf0f1;
            border-radius: 5px;
            flex: 1;
            margin: 0 10px;
        }}
        .stat-number {{
            font-size: 2em;
            font-weight: bold;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>Отчёт анализа документа</h1>
        <p>Документ: {self.report.document_name}</p>
        <p>Дата: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
    </div>
    
    <div class="section">
        <h2>Статистика</h2>
        <div class="stats">
            <div class="stat-box">
                <div class="stat-number">{self.report.total_chunks}</div>
                <div>Частей</div>
            </div>
            <div class="stat-box">
                <div class="stat-number">{self.report.total_findings}</div>
                <div>Проблем</div>
            </div>
            <div class="stat-box">
                <div class="stat-number severity-high">{high}</div>
                <div>Высокая</div>
            </div>
            <div class="stat-box">
                <div class="stat-number severity-medium">{medium}</div>
                <div>Средняя</div>
            </div>
            <div class="stat-box">
                <div class="stat-number severity-low">{low}</div>
                <div>Низкая</div>
            </div>
        </div>
    </div>
"""
        
        # Global analysis
        if self.report.global_analysis:
            html += f"""
    <div class="section">
        <h2>Глобальный анализ</h2>
        <h3>Общее резюме</h3>
        <p>{self.report.global_analysis.overall_summary}</p>
"""
            
            if self.report.global_analysis.document_structure_issues:
                html += "<h3>Проблемы структуры</h3><ul>"
                for issue in self.report.global_analysis.document_structure_issues:
                    html += f"<li>{issue}</li>"
                html += "</ul>"
            
            if self.report.global_analysis.consistency_issues:
                html += "<h3>Проблемы согласованности</h3><ul>"
                for issue in self.report.global_analysis.consistency_issues:
                    html += f"<li>{issue}</li>"
                html += "</ul>"
            
            if self.report.global_analysis.recommendations:
                html += "<h3>Рекомендации</h3><ul>"
                for rec in self.report.global_analysis.recommendations:
                    html += f"<li>{rec}</li>"
                html += "</ul>"
            
            html += "</div>"
        
        # Detailed findings
        html += '<div class="section"><h2>Детальные результаты</h2>'
        
        for chunk in self.report.chunk_analyses:
            html += f'<h3>Часть {chunk.chunk_id}</h3>'
            html += f'<p><strong>Краткое содержание:</strong> {chunk.summary}</p>'
            
            if chunk.findings:
                for finding in chunk.findings:
                    severity_class = f"severity-{finding.severity}"
                    html += f'''
                    <div class="finding">
                        <p><span class="{severity_class}">[{finding.severity.upper()}]</span> {finding.issue_type}</p>
                        <div class="quote">"{finding.quote}"</div>
                        <p><strong>Проблема:</strong> {finding.problem}</p>
                        <p><strong>Рекомендация:</strong> {finding.recommendation}</p>
                    </div>
                    '''
        
        html += """
    </div>
</body>
</html>
"""
        return html
